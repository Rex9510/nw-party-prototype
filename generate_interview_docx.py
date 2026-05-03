# -*- coding: utf-8 -*-
"""Generate 党务智能体建设需求访谈确认表.docx"""
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROWS = [
    ("范围与分期", "首期上线必须包含哪些功能？哪些明确放到二期、三期？请按「必须有 / 最好有 / 暂不做」排序。", "【业务】"),
    ("范围与分期", "是否接受分期验收？一期「先能用」的业务标准是什么（例如先解决哪几类党员办事、哪几类党务问答）？", "【业务】"),
    ("范围与分期", "本期「明确不做」的功能清单能否先写一版，避免实施过程中范围蔓延？", "【业务】"),
    ("范围与分期", "用户入口形态：小程序、H5、PC 管理端分别是否需要？是否与现有公众号/街道统一入口做同一套体验？", "【业务】"),
    ("范围与分期", "服务对象范围：仅本街道党员与党务工作者，还是包含流动党员报到对象、群众咨询等？", "【业务】"),
    ("使用规模与场景", "本街道党员总数、流动党员大致占比；首期希望覆盖到多少人「愿意常用」？", "【业务】→【业务→技术】"),
    ("使用规模与场景", "党务智能问答：希望解决哪几类高频问题（请列 10～20 个典型问法）。", "【业务】"),
    ("使用规模与场景", "是否会出现「上传长文档/PDF 让系统解读」的场景？多不多？", "【业务→技术】"),
    ("使用规模与场景", "三会一课等台账：每次需上传几张照片、附件大致多大；各社区大概频率？", "【业务】→【业务→技术】"),
    ("使用规模与场景", "党校直播：一年大概多少场、单场希望多少人同时在线（量级即可）？是否需要长期回看？", "【业务】→【业务→技术】"),
    ("档案与画像", "党员档案在业务上主要是结构化信息，还是大量扫描件/复印件需要检索与核对？", "【业务】"),
    ("档案与画像", "「档案智能审核」希望产出物是什么：问题清单、缺项提示、修改建议，还是还要自动生成某种文书？最终是否必须人工确认？", "【业务】"),
    ("档案与画像", "不合格党员相关功能：是否约定为「政策依据 + 流程模板 + 材料参考」，禁止系统自动作出组织处理结论？（建议组织/法务同框）", "【业务】+【技术/政务/合规】"),
    ("档案与画像", "党员画像：最想从哪些工作体现（学习、活动、志愿服务、履职等）？画像主要给谁看、用于什么管理场景？", "【业务】"),
    ("档案与画像", "除姓名、手机、住址外，系统里是否还有身份证、奖惩材料全文等？哪些字段绝对不能对外展示或不能进入智能问答上下文？（字段业务列，规则合规定）", "【业务】+【技术/政务/合规】"),
    ("档案与画像", "党员转出、离职后数据是否保留、保留多久、能否导出删除——业务管理制度上的要求是什么？", "【业务】+【技术/政务/合规】"),
    ("对接与协同", "对接共产党员网、学习强国、深圳智慧党建：业务上必须达到哪种效果——跳转学习即可、能登记学习记录即可，还是必须自动同步学习数据？", "【业务】→【技术/政务】"),
    ("对接与协同", "对接「深小服」：希望用户在本系统里完成哪些具体办事体验（请列 3～5 个例子）。", "【业务】→【技术/政务】"),
    ("对接与协同", "登录账号：希望党员「少记一套密码」与市/区统一身份打通，还是街道自建账号也可接受？", "【业务】→【技术/政务】"),
    ("对接与协同", "是否需要短信验证、实名认证、人脸等高强度身份手段？业务上是否强制？", "【业务】→【技术/采购】"),
    ("对接与协同", "除方案已列系统外，是否还必须与其他系统打通？请列「希望打通的系统 + 用来解决什么问题」。", "【业务】→【技术】"),
    ("大模型与问答体验", "党务问答验收口径：是否要求「回答必须引用依据、禁止编造」？能否接受「不知道就明确说不知道并引导人工」？", "【业务】"),
    ("大模型与问答体验", "知识库材料：首版由哪些科室提供、更新节奏希望怎样（如每季度）？谁签字对内容负责？", "【业务】"),
    ("大模型与问答体验", "调用云上模型时，问答日志是否允许用于厂商模型改进——单位是否有硬性禁止？", "【技术/政务/合规】"),
    ("安全与等保", "系统计划部署在互联网区、政务外网还是指定政务云？是否有指定云平台？", "【技术/政务/采购】"),
    ("安全与等保", "主管单位是否已要求必须做网络安全等级保护？目标级别是否有指示？", "【技术/政务/合规】"),
    ("安全与等保", "是否还有商用密码应用安全性评估等单独要求？", "【技术/政务/合规】"),
    ("安全与等保", "操作与问答日志：希望留存多久、哪些角色可审计——业务提管理需求即可。", "【业务】→【技术/政务/合规】"),
    ("安全与等保", "运维是否允许远程登录生产、是否需要驻场——有无硬性规定？", "【技术/政务/合规/采购】"),
    ("党校与培训", "课程中心：以外链聚合为主，还是必须支持本街道上传课程与版权自持？", "【业务】"),
    ("党校与培训", "直播：更倾向「够用即可的第三方 SaaS」还是「必须统一纳入某平台」？", "【业务】→【技术/采购】"),
    ("党校与培训", "参训记录是否与评优、画像指标挂钩？统计口径谁定？", "【业务】"),
    ("运营与推广", "上线后知识库更新、热点问题维护、敏感答复审核，业务侧准备安排哪些岗位、大约多少人力？", "【业务】"),
    ("运营与推广", "各社区党务干部培训与推广是否有统一时间表或考核节点？", "【业务】"),
    ("采购与验收", "采购方式与流程大致需要多长时间（是否已立项、预算是否已批）？", "【采购/管理】"),
    ("采购与验收", "合同验收标准：除功能清单外，是否规定性能、安全、等保测评通过等为验收前置？（建议联合定稿）", "【业务】+【技术/政务/合规/采购】"),
    ("采购与验收", "是否有必须上线的业务 deadline（如某主题活动前）？", "【业务/管理】"),
]


def set_run_font(run, name="微软雅黑", size_pt=10.5):
    run.font.name = name
    run.font.size = Pt(size_pt)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def main():
    # English filename avoids mojibake on some Windows shells; title inside doc is Chinese.
    out = Path(__file__).resolve().parent / "party_affairs_interview_confirmation.docx"
    doc = Document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("南湾街道党员党务智能体建设需求访谈确认表")
    set_run_font(r, "微软雅黑", 16)
    r.bold = True

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run("（用于需求确认后再评估费用与建设周期）")
    set_run_font(r2, "微软雅黑", 10)

    doc.add_paragraph()

    headers = ["序号", "模块", "访谈问题", "责任标签", "业务方确认结果", "备注"]
    table = doc.add_table(rows=1 + len(ROWS), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    for i, text in enumerate(headers):
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_run_font(run, "微软雅黑", 10)
        run.bold = True

    for idx, (mod, question, tag) in enumerate(ROWS, start=1):
        row = table.rows[idx].cells
        cells = [str(idx), mod, question, tag, "", ""]
        for col_idx, text in enumerate(cells):
            p = row[col_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if col_idx == 2 else (
                WD_ALIGN_PARAGRAPH.CENTER if col_idx in (0, 1) else WD_ALIGN_PARAGRAPH.LEFT
            )
            run = p.add_run(text)
            set_run_font(run, "微软雅黑", 10)

    # Column widths (total ~17cm for A4 portrait with margins)
    widths = [Cm(1.0), Cm(2.2), Cm(7.8), Cm(3.0), Cm(2.5), Cm(2.0)]
    for row in table.rows:
        for c_idx, cell in enumerate(row.cells):
            cell.width = widths[c_idx]

    doc.save(out)
    print("Wrote:", out)


if __name__ == "__main__":
    main()
